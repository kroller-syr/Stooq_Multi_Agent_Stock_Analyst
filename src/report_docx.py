import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


def save_report_docx(
    ticker: str,
    report_markdown: str,
    references_text: str,
    figure_paths: list[str],
    reports_dir: str = "outputs/reports",
) -> str:
    os.makedirs(reports_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = Path(reports_dir) / f"{ticker}_report_{ts}.docx"

    doc = Document()
    doc.add_heading(f"{ticker} — Stock & News Analysis Report", level=1)

    # Body: treat markdown-ish text as paragraphs (simple v1)
    for line in (report_markdown or "").splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            # markdown headers -> doc headings (rough)
            level = min(4, max(1, line.count("#")))
            doc.add_heading(line.lstrip("#").strip(), level=level)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    # Figures
    if figure_paths:
        doc.add_page_break()
        doc.add_heading("Figures", level=1)
        for idx, p in enumerate(figure_paths, start=1):
            pth = Path(p)
            if pth.exists():
                doc.add_paragraph(f"Figure {idx}: {pth.name}")
                doc.add_picture(str(pth), width=Inches(6.5))
            else:
                doc.add_paragraph(f"Figure {idx}: (missing) {pth.name}")

    # References
    if references_text:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        for line in references_text.splitlines():
            doc.add_paragraph(line)

    doc.save(out_path)
    return str(out_path)